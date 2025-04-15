from docx import Document
from lxml import etree
import json
from docx.shared import Pt, RGBColor
from collections import Counter
from enum import Enum
import pypandoc
import re
from bs4 import BeautifulSoup
import os
import zipfile
import shutil
import xml.etree.ElementTree as ET
import html  # Ensure this is imported

# Define the prefix for alt text that should be kept in the output
ALT_TEXT_KEEP_PREFIX = "keep-"

## Define the folder structure for the output
FOLDERS = {
    "content": "content",
    "media": "assets",
    "data": "data",
    "js": "js",
    "css": "css",
}

# Define system-wide defaults as a simple dictionary
GLOBAL_DEFAULT_STYLES = {
    "textAlign": "left",
    "fontSize": "1rem",  # 16px converted to rem
    "fontWeight": "normal",
    "color": "#000000",
    "backgroundColor": "#FFFFFF",
    "padding": "8px",  # 5px converted to rem
    "verticalAlign": "middle"
}

print(GLOBAL_DEFAULT_STYLES)

def check_compatibility(docx_path, output_path):
    """Check compatibility of docx file and fix or flag any potential issues."""
    ## Load the document
    doc = Document(docx_path)
    ## normalize empty paragraphs in a DOCX file
    def normalize_empty_paragraphs(doc):      
        for para in doc.paragraphs:
            # Check if the paragraph is empty but has a non-Normal style
            if not para.text.strip() and para.style.name != "Normal":
                para.style = doc.styles["Normal"]
        return doc

    doc = normalize_empty_paragraphs(doc)
    # Save the modified document
    doc.save(output_path)
    print(f"Compatible document saved as: {output_path}")

import zipfile
import xml.etree.ElementTree as ET


def check_for_missing_figures(docx_path, html_content):
    """
    Check for missing figures in the DOCX file compared to the HTML content.
    This function extracts figures from the DOCX file and verifies if they are present in the HTML.

    Args:
        docx_path (str): Path to the DOCX file.
        html_content (str): HTML content as a string.

    Returns:
        list: List of missing figures in the HTML content.
    """
    # Namespaces used in Word XML
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    }

    figures = []

    with zipfile.ZipFile(docx_path) as docx:
        # Step 1: Extract only the document.xml
        with docx.open('word/document.xml') as file:
            tree = ET.parse(file)
            root = tree.getroot()

        # Step 2: Extract image and caption info
        paragraphs = root.findall('.//w:p', ns)

        for i, p in enumerate(paragraphs):
            img_tag = p.find('.//a:blip', ns)
            if img_tag is not None:
                r_id = img_tag.attrib.get(f'{{{ns["r"]}}}embed')
                # Look ahead for the caption (next paragraph)
                caption_text = ''
                if i + 1 < len(paragraphs):
                    next_p = paragraphs[i + 1]
                    caption_text = ''.join([t.text or '' for t in next_p.findall('.//w:t', ns)])
                figures.append({'rId': r_id, 'caption': caption_text})

    # Step 3: Verify which figures have captions starting with "Figure {d} or Figure {letters}"
    figures_verified = []
    for fig in figures:
        caption = fig['caption']
        if caption.startswith("Figure"):
            match = re.search(r"Figure [A-Za-z0-9]+", caption)
            if match:
                figure_number = match.group(0)
                caption_after = caption.split(figure_number)[1].strip()
                figures_verified.append({'rId': fig['rId'], 'caption': caption, 'html_caption': caption_after})

    # Step 4: Check for missing figures in the html
    missing_figures = []
    soup = BeautifulSoup(html_content, 'html.parser')
    figure_tags = soup.find_all('figcaption')

    # Identify missing captions that are in the docx but not in the html
    for fig in figures_verified:
        found = False
        for fig_tag in figure_tags:
            if fig['html_caption'] in fig_tag.text.replace("\r\n", " "):
                found = True
                break
        if not found:
            missing_figures.append(fig['caption'])

    return missing_figures

# Function to map Word vertical alignment values to CSS
def map_vertical_align(word_val: str) -> str:
    """Convert Word's vertical alignment values to CSS equivalents."""
    mapping = {
        "top": "top",
        "center": "middle",  # Word uses "center", CSS uses "middle"
        "bottom": "bottom",
        "both": "middle"  # No direct equivalent, "middle" is a safe choice
    }
    return mapping.get(word_val.lower(), "middle")  # Default to "middle"


def convert_pt_to_rem(pt_size, base_font_size=16):
    """
    Converts font sizes from a DOCX file (in pt) to rem for HTML rendering.
    :param pt_size: Font size in points.
    :param base_font_size: Base font size in pixels (default: 16px).
    :return: Dictionary mapping text to its font size in rem.
    """
    rem_size = (pt_size * 1.3333) / base_font_size

    return f"{rem_size:.2f}rem"

def get_font_info(style):
    """Extract font information from a paragraph or run style."""
    font_info = {}

    if style and style.font:
        if style.font.name:
            font_info["fontFamily"] = style.font.name
        if style.font.size:
            font_info["fontSize"] = convert_pt_to_rem(style.font.size.pt)
        if style.font.bold:
            font_info["fontWeight"] = "bold"
        if style.font.italic:
            font_info["fontStyle"] = "italic"
        if style.font.color and style.font.color.rgb:
            font_info["color"] = f"#{style.font.color.rgb}"  # Convert RGBColor to hex

    return font_info

## TABLES
# Default styles for tables
DEFAULT_STYLES = {
        "table": {

        },
        "th": {

        },
        "td1": {

        },
        "td": {

        }
    }


def get_table_alignment(table):
    tbl_pr = table._element.xpath(".//w:jc")
    if tbl_pr:
        return tbl_pr[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
    return "default"

def get_cell_vertical_alignment(cell):
    v_align = cell._element.xpath(".//w:vAlign")
    if v_align:
        return v_align[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
    return "center"

def get_paragraph_alignment(paragraph):
    p_pr = paragraph._element.xpath(".//w:jc")
    if p_pr:
        return p_pr[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
    return "left"

def clean_dict(data):
    """
    Recursively removes keys with values of '#auto' or None from a dictionary.
    """
    if isinstance(data, dict):
        return {k: clean_dict(v) for k, v in data.items() if v not in ('#auto', None)}
    elif isinstance(data, list):
        return [clean_dict(v) for v in data]
    else:
        return data


# Function to compute the modal values while ignoring missing properties
def compute_modal_style(styles_list):
    # Add the default styles to ensure they are considered in the modal calculation
    styles_list.append(GLOBAL_DEFAULT_STYLES.copy())

    modal_styles = {}
    all_keys = {key for styles in styles_list for key in styles}
    
    for key in all_keys:
        values = [styles[key] for styles in styles_list if key in styles]
        if values and Counter(values).most_common(1)[0][0] is not None:
            modal_styles[key] = Counter(values).most_common(1)[0][0]

    return modal_styles

def extract_table_default_styles(doc) -> dict:
    table_styles = {
        "table": {
            "border": "1px solid black",
            "borderCollapse": "collapse",
            "marginBottom": "12px"
        },
        "th": {},
        "td1": {},
        "td": {}
    }

    if not doc.tables:
        return table_styles

    table = doc.tables[1]  # Assume second table defines the styles
    default_font_size_rem = convert_pt_to_rem(get_doc_default_font_size(doc)) ## Default font size extraction

    th_styles = []
    td1_styles = []
    td_styles = []


    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_style = {}

            # Extract text alignment
            text_align = None
            for para in cell.paragraphs:
                if para.text.strip():
                    text_align = get_paragraph_alignment(para)
                    break
            if text_align:
                cell_style["textAlign"] = text_align

            # Extract font styles
            font_size = None
            font_weight = None
            font_color = None

            for para in cell.paragraphs:
                if para.text.strip():
                    for run in para.runs:
                        if run.font.size:
                            font_size = convert_pt_to_rem(run.font.size.pt)
                        elif para.style and para.style.font.size:
                            font_size = convert_pt_to_rem(para.style.font.size.pt)

                        if run.bold:
                            font_weight = "bold"

                        if run.font.color and run.font.color.rgb:
                            font_color = f"#{run.font.color.rgb}"
                        elif para.style and para.style.font.color and para.style.font.color.rgb:
                            font_color = f"#{para.style.font.color.rgb}"
                    break  # Only need the first valid paragraph with text

            # if font_size:
            #     cell_style["fontSize"] = font_size
            # if font_weight:
            #     cell_style["fontWeight"] = font_weight
            # if font_color:
            #     cell_style["color"] = font_color
            cell_style["fontSize"] = font_size if font_size else default_font_size_rem
            cell_style["fontWeight"] = font_weight if font_weight else None
            cell_style["color"] = font_color if font_color else None

            # Extract background color
            shading = cell._element.xpath('.//w:shd/@w:fill')
            cell_style["backgroundColor"] = f"#{shading[0]}" if shading else None  # Leave None to filter later
            
            # Padding and vertical alignment
            cell_style["padding"] = "5px"
            cell_style["verticalAlign"] = map_vertical_align(get_cell_vertical_alignment(cell))

            # Store styles in respective lists
            if row_idx == 0:
                th_styles.append(cell_style)
            elif col_idx == 0:
                td1_styles.append(cell_style)
            else:
                td_styles.append(cell_style)

    # Compute the modal styles for each category
    table_styles["th"] = compute_modal_style(th_styles)
    table_styles["td1"] = compute_modal_style(td1_styles)
    table_styles["td"] = compute_modal_style(td_styles)

    return clean_dict(table_styles)


## END TABLES

## Default font size extraction
def get_style_font_size(doc, style_name):
    """Fetch font size from a named style if available."""
    try:
        style = doc.styles[style_name]
        if style and style.font and style.font.size:
            return style.font.size.pt
    except KeyError:
        pass
    return None  # Style not found or no explicit size set



def get_doc_default_font_size(doc):
    """Extract the default font size from the document."""
    # Get the XML element as a string
    xml_data = doc.styles.element.xml

    # Define the XML namespace
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Parse the XML
    root = ET.fromstring(xml_data)

    # Find the <w:docDefaults> section and locate <w:sz>
    sz_element = root.find('.//w:docDefaults//w:sz', ns)

    # Extract the font size
    if sz_element is not None:
        font_size = sz_element.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        return int(font_size) / 2 # Convert half-points to points 
    return None


def extract_styles(doc_path):
    """Extract styles from the Word document dynamically."""
    doc = Document(doc_path)

    styles_data = {
        "headings": {},
        "body": {},
        "lists": {},
        "captions": {},
    }

    for para in doc.paragraphs:
        para_style = para.style

        if para_style and para_style.name.startswith("Heading"):
            heading_level = para_style.name.replace("Heading ", "h")
            styles_data["headings"][heading_level] = get_font_info(para_style)

        elif para_style and "Caption" in para_style.name:
            styles_data["captions"]["caption"] = get_font_info(para_style)

        elif para_style and "List" in para_style.name:
            list_type = "ul" if "Bullet" in para_style.name else "ol"
            styles_data["lists"][list_type] = get_font_info(para_style)

        elif para_style and para_style.name == "Normal":
            styles_data["body"]["p"] = get_font_info(para_style)

    # Extract default font size
    default_font_size = get_style_font_size(doc, "Normal") or get_doc_default_font_size(doc) or 12
    styles_data["body"]["p"]["fontSize"] = convert_pt_to_rem(default_font_size)

    ## TABLES
    if doc.tables:
        table = extract_table_default_styles(doc)
        styles_data.update(table)
    return styles_data

def is_cell_merged(cell, row_idx, col_idx, merge_tracker):
    """Check if a Word table cell is merged (horizontally or vertically) and determine row span."""
    cell_xml = cell._tc  # Get the XML element of the table cell

    # Check for horizontal merge (gridSpan)
    grid_span = cell_xml.xpath('.//w:gridSpan')
    col_span = int(grid_span[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", 1)) if grid_span else 1

    # Check for vertical merge (vMerge)
    v_merge = cell_xml.xpath('.//w:vMerge')
    row_span = None  # Only set if it's actually merged

    if v_merge:
        v_merge_val = v_merge[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
        if v_merge_val == "restart":  # Start of a vertically merged section
            merge_tracker[(row_idx, col_idx)] = 1  # Initialize tracking
        elif v_merge_val is None:  # Continuation of merge
            for r in range(row_idx - 1, -1, -1):  # Find the start of this vertical merge
                if (r, col_idx) in merge_tracker:
                    merge_tracker[(r, col_idx)] += 1
                    row_span = merge_tracker[(r, col_idx)]
                    return {"hidden": True}  # Mark this cell as part of the merged block, but not stored separately

    return {
        "colSpan": col_span if col_span > 1 else None,  # Store only if >1
        "rowSpan": row_span if row_span and row_span > 1 else None  # Store only if >1
    }




from docx import Document

def extract_table_format(doc_path, default_styles: dict = DEFAULT_STYLES):
    doc = Document(doc_path)
    tables_info = {}

    for table_idx, table in enumerate(doc.tables):
        print(f"Processing table {table_idx + 1}...")
        table_id = f"table_{table_idx}"  # Generate table ID
        table_info = {
            "headers": [],
            "rows": []
        }

        merge_tracker = {}  # Track merged cells {(row_idx, col_idx): remaining_span}

        for row_idx, row in enumerate(table.rows):
            row_data = []
            col_idx = 0  # Track actual column position for skipping merged cells

            while col_idx < len(row.cells):
                if (row_idx, col_idx) in merge_tracker:
                    merge_tracker[(row_idx, col_idx)] -= 1
                    if merge_tracker[(row_idx, col_idx)] == 0:
                        del merge_tracker[(row_idx, col_idx)]  # Clear when done
                    col_idx += 1
                    continue  # Skip merged cells

                cell = row.cells[col_idx]

                # Determine default style based on position
                if row_idx == 0:
                    default_style = default_styles["th"]
                elif col_idx == 0:
                    default_style = default_styles["td1"]
                else:
                    default_style = default_styles["td"]

                # Extract formatting and text as separate parts
                text_parts = []
                prev_para = None  # Track previous paragraph for detecting paragraph breaks

                for para in cell.paragraphs:
                    # Extract alt text from images in the paragraph, if the paragraph has an element
                    img_matches = []
                    if para._element is not None:
                        for drawing in para._element.findall('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            print("Found a drawing element")
                            doc_pr = drawing.find('.//wp:docPr', namespaces={'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                            if doc_pr is not None:
                                alt_text = doc_pr.attrib.get('descr', '').strip()
                                if alt_text:
                                    img_match = {"alt_text": (alt_text)}
                                    _, image_type, *image_alt_text = alt_text.split("-")
                                    image_alt_text = "-".join(image_alt_text)
                                    icon_class, icon_name = image_alt_text.split(":")
                                    if image_type == "icon":
                                        img_match["iconHtml"] = f"<span class={icon_class}>{icon_name}</span>"
                                        img_match["iconPosition"] = "start" #TODO make dynamic based on actual position in cell: could also be top, bottom, end
                                    else:
                                        print("Images not supported in tables yet...")
                            img_matches.append(img_match)

                    if para.text.strip() or any(run.text.strip() for run in para.runs):  # Ignore empty paragraphs
                        # Insert newline if this is a new paragraph (except for the first one)
                        if prev_para is not None:
                            text_parts.append({"text": "\n", "newline": True})

                        prev_para = para  # Update previous paragraph tracker

                        for run in para.runs:
                            if run.text:
                                # Handle inline newlines within a run
                                segments = run.text.split("\n")
                                for i, segment in enumerate(segments):
                                    part = {"text": segment}
                                    if run.bold:
                                        part["bold"] = True
                                    if run.italic:
                                        part["italic"] = True
                                    if run.font.superscript:
                                        part["superscript"] = True
                                    if run.font.subscript:
                                        part["subscript"] = True
                                    if run.font.color and run.font.color.rgb:
                                        part["color"] = f"#{run.font.color.rgb}" if f"#{run.font.color.rgb}" != default_style.get("color") else None
                                    if run.font.size:
                                        part["fontSize"] = convert_pt_to_rem(run.font.size.pt)

                                    text_parts.append(part)

                                    # If this was a split part, add an explicit newline
                                    if i < len(segments) - 1:
                                        text_parts.append({"text": "\n", "newline": True})

                # Check if any part has formatting
                has_formatting = any(len(part) > 1 for part in text_parts if part["text"] != "\n")

                # Check if all formatted parts share the same styling
                def extract_format(part):
                    """Extracts formatting keys (excluding text) from a part."""
                    return {k: v for k, v in part.items() if k not in ["text", "newline"]}

                common_format = extract_format(text_parts[0]) if text_parts else {}
                all_same_format = all(extract_format(part) == common_format for part in text_parts if part["text"] != "\n")

                # Handle different cases
                if not text_parts:  # No text, return empty string
                    actual_style = {"text": ""}
                elif not has_formatting:  # Merge unformatted text, preserving newlines
                    actual_style = {
                        "text": "".join(
                            part["text"] if "newline" not in part else "<br>"
                            for part in text_parts
                        )
                    }
                elif all_same_format:  # Merge text and apply common formatting at the top level, preserving newlines
                    actual_style = {
                        "text": "".join(
                            part["text"] if "newline" not in part else "<br>"
                            for part in text_parts
                        ),
                        **common_format  # Apply the shared formatting
                    }
                else:  # Mixed formatting, keep textParts
                    actual_style = {"textParts": text_parts}


                # Handle merged cells
                merge_info = is_cell_merged(cell, row_idx, col_idx, merge_tracker)
                if "hidden" in merge_info:
                    col_idx += 1
                    continue  # Skip storing this cell (it's a continuation of a merged cell)

                if merge_info["colSpan"]:
                    actual_style["colSpan"] = merge_info["colSpan"]
                    for i in range(1, merge_info["colSpan"]):  # Skip following columns
                        merge_tracker[(row_idx, col_idx + i)] = 1

                if merge_info["rowSpan"]:
                    actual_style["rowSpan"] = merge_info["rowSpan"]
                    for i in range(1, merge_info["rowSpan"]):  # Track vertically merged cells
                        merge_tracker[(row_idx + i, col_idx)] = merge_info["rowSpan"] - i

                # Extract other styles and replace if they differ from default
                cell_text_align = get_paragraph_alignment(cell.paragraphs[0]) if cell.paragraphs else default_style.get("textAlign", "left")
                
                actual_style["textAlign"] = cell_text_align if cell_text_align != default_style["textAlign"] else None
                cell_vertical_align = map_vertical_align(get_cell_vertical_alignment(cell))
                actual_style["verticalAlign"] = cell_vertical_align if cell_vertical_align != default_style["verticalAlign"] else None

                # Extract background color
                shading = cell._element.xpath('.//w:shd/@w:fill')
                actual_style["backgroundColor"] = f"#{shading[0]}" if shading else None

                if img_matches:
                    for img in img_matches: ## will be an empty list if none were found
                        if "iconHtml" in img:
                            actual_style["iconHtml"] = img["iconHtml"]
                            actual_style["iconPosition"] = img["iconPosition"]
                    

                # Remove default styles
                actual_style = clean_dict(actual_style)

                # Determine if this is a header row or data row
                if row_idx == 0:
                    table_info["headers"].append(actual_style)
                else:
                    row_data.append(actual_style)

                col_idx += 1  # Move to the next column

            if row_idx > 0:  # Store only data rows (headers handled separately)
                table_info["rows"].append(row_data)

        tables_info[table_id] = table_info

    return tables_info




def extract_docx(doc_path, output_path):
    """Extracts the DOCX contents into a specified folder."""
    extracted_folder = os.path.join(output_path, "docx_extracted")
    os.makedirs(extracted_folder, exist_ok=True)
    
    with zipfile.ZipFile(doc_path, "r") as docx_zip:
        docx_zip.extractall(extracted_folder)
    
    return extracted_folder

def parse_relationships(rels_path):
    """Parses the relationships file to map image IDs to filenames."""
    image_map = {}
    
    if os.path.exists(rels_path):
        tree = ET.parse(rels_path)
        root = tree.getroot()
        
        for rel in root.findall(".//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"):
            rid = rel.attrib.get("Id", "")
            target = rel.attrib.get("Target", "")
            
            if "media/" in target:
                image_map[rid] = target.split("/")[-1]
    
    return image_map

def extract_alt_texts(doc_xml_path, image_map, allowed_alt_texts, extracted_folder, output_path, image_folder):
    """Extracts images based on allowed alt texts and renames them."""
    alt_text_map = {}

    # Create the output folder for images
    media_folder = os.path.join(output_path, image_folder)
    print(f"📂 Creating media folder: {media_folder}")
    os.makedirs(image_folder, exist_ok=True)
    
    if os.path.exists(doc_xml_path):
        tree = ET.parse(doc_xml_path)
        root = tree.getroot()

        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main", 
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        }

        for drawing in root.findall(".//w:drawing", ns):
            doc_pr = drawing.find(".//a:blip", ns)
            descr_tag = drawing.find(".//wp:docPr", ns)
            
            if doc_pr is not None and descr_tag is not None:
                alt_text = descr_tag.attrib.get("descr", "").strip()
                rid = doc_pr.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed", "")
                
                if rid in image_map and (alt_text in allowed_alt_texts or alt_text.startswith(ALT_TEXT_KEEP_PREFIX)):
                    old_name = image_map[rid]
                    alt_text_map[os.path.splitext(old_name)[0]] = f"{image_folder}/{old_name}"

                    old_path = os.path.join(extracted_folder, "word/media", old_name)
                    new_path = os.path.join(media_folder, old_name)

                    if os.path.exists(old_path):
                        print(f"✅ Moving image: {old_name} ➝ {new_path}")
                        shutil.move(old_path, new_path)
                    else:
                        print(f"❌ ERROR: Image file not found: {old_path}")
    
    return alt_text_map


def extract_docx_media(doc_path, output_path, media_folder, allowed_alt_texts):
    """Extracts DOCX contents and images based on allowed alt texts."""
    # Extract DOCX contents
    extracted_folder = extract_docx(doc_path, output_path)

    # Paths to XML files inside the extracted DOCX
    rels_path = os.path.join(extracted_folder, "word/_rels/document.xml.rels")
    doc_xml_path = os.path.join(extracted_folder, "word/document.xml")

    # Parse relationships to map image IDs to filenames
    image_map = parse_relationships(rels_path)

    alt_text_map = extract_alt_texts(doc_xml_path, image_map, allowed_alt_texts, extracted_folder, output_path, media_folder)

    return alt_text_map

def generate_lua_lookup_table(image_numbers):
    """
    Generates a Pandoc-compatible metadata JSON for Lua.
    Example: [3, 5, 6] -> { "keep_images": [3, 5, 6] }
    """
    metadata = {
        "keep_images": image_numbers  # Ensure it's a list, not a string
    }

    # Convert metadata to JSON
    metadata_json = json.dumps(metadata)

    return metadata_json

def remove_empty_figures(html):
    # Regex pattern to match <figure> tags that do not contain an <img>
    pattern = re.compile(r'<figure>(?:(?!<img).)*?</figure>', re.DOTALL)

    # Remove matching <figure> elements
    cleaned_html = re.sub(pattern, '', html)

    return cleaned_html


def remove_empty_paragraphs(html):
    if not isinstance(html, str):  # Ensure input is a string
        raise TypeError(f"Expected a string, but got {type(html).__name__}")

    pattern = re.compile(r'<p>\s*<em>\s*<br\s*/?>\s*\r?\n?\s*</em>\s*</p>', re.IGNORECASE)
    
    # Remove empty paragraphs
    cleaned_html = re.sub(pattern, '', html)

    return cleaned_html

# Function to replace tables and store captions safely
def table_replacer(match, counter=[0], table_caption_counter=[0]):
    table_html = match.group(0)
    table_number = counter[0]
    # Extract caption if present
    caption_match = re.search(r'<caption.*?>(.*?)</caption>', table_html, re.DOTALL)
    if caption_match:
        caption_text = caption_match.group(1).strip()
    else:
        counter[0] += 1
        return f'<div class="table" id="table_{table_number}"></div><br>'

    # Add in caption prefix that is misplaced by pypandoc
    if not caption_text.startswith("<p>Table"):
        if caption_text.startswith("<p>"):
            caption_text = caption_text[3:]
        print(f"'Table ' not found in caption text for table {table_number}. Adding manually as: Table {table_caption_counter[0]+1}{caption_text}"[:-4])
        caption_text = f"Table {table_caption_counter[0]+1}{caption_text}"[:-4] ## [:-4] to remove </p> tag

    # Escape caption for safe inclusion in HTML attributes
    caption_escaped = html.escape(caption_text)  # <-- This should work now

    # Replace table with a div containing the caption as a data attribute
    replacement = f'<div class="table" id="table_{table_number}" data-caption="{caption_escaped}"></div><br>'
    
    counter[0] += 1
    table_caption_counter[0] += 1

    return replacement


def generate_navigation_data(html_content):
    """
    Parses HTML content to generate structured navigation data.

    :param html_content: HTML content as a string.
    :return: Structured navigation data as a list of dictionaries.
    """
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Organize headings into a structured hierarchy
    nav_data = []
    current_h1 = None
    current_h2 = None

    for tag in soup.find_all(['h1', 'h2', 'h3']):
        heading_id = tag.get('id')
        if not heading_id:
            heading_id = re.sub(r'\s+', '-', tag.text.strip().lower())
            tag['id'] = heading_id

        if tag.name == 'h1':
            current_h1 = {
                "id": heading_id,
                "text": tag.text,
                "h2": []
            }
            nav_data.append(current_h1)

        elif tag.name == 'h2' and current_h1:
            current_h2 = {
                "id": heading_id,
                "text": tag.text,
                "h3": []
            }
            current_h1["h2"].append(current_h2)

        elif tag.name == 'h3' and current_h1 and current_h2:
            current_h2["h3"].append({
                "id": heading_id,
                "text": tag.text
            })

    return nav_data

def insert_sub_navigation(html_content, nav_json):
    """
    Inserts <div data-sub-navigation data-parent="section_id"></div> before the first <h2> of each section.
    """
    for section in nav_json:
        section_id = section["id"]
        h2_items = section.get("h2", [])
        
        # Only proceed if there are H2 elements in the section
        if not h2_items:
            continue
        
        # Define the sub-navigation div to be inserted
        sub_nav_div = f'\n<div data-sub-navigation data-parent="{section_id}"></div>'
        
        # Find the first H2 in this section and insert the sub-navigation div before it
        pattern = rf'(\n<h2 id="{h2_items[0]["id"]}")'
        html_content = re.sub(pattern, sub_nav_div + r'\1', html_content, count=1)
    
    return html_content



## Image html replacement
import re
import os
import html

def find_image_alt_text(html_content):
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a mapping of image filename to alt text from the HTML
    html_alt_map = {
        img['src']: img.get('alt', '')
        for img in soup.find_all('img') if 'src' in img.attrs
    }

    return html_alt_map


from collections import namedtuple

## figure number
def get_figure_captions(html_content, doc_img_src: list) -> namedtuple:
    soup = BeautifulSoup(html_content, 'html.parser')
    figure_tags = soup.find_all('figure')
    figure_number = 0
    figure_number_new = 0
    Image = namedtuple('Image', ["figure_number", "figure_number_new", "figure_caption", "alt_text"])
    result = {}

    for figure in figure_tags:
        figure_number += 1
        if not figure.find('img') or figure.find('img')['src'] not in doc_img_src:
            continue
        figure_number_new += 1
        figure_caption = figure.find('figcaption').text if figure.find('figcaption') else ""
        alt_text = figure.find('img')['alt'] if figure.find('img')['alt'] else ""
        current_img = os.path.splitext(os.path.basename(figure.find('img')['src']))[0]
        result[current_img] = Image(figure_number, figure_number_new, figure_caption, alt_text)

    return result

# def identify_image_type(images_dict: dict):
#     # split into icons, charts, and images
#     icons = {}
#     charts = {}
#     images = {}
#     for key, value in images_dict.items():
#         print(value)
#         if value.startswith("icon"):
#             print(f"starts with icon: {value}")
#             icons[key] = value.replace("icon-", "")
#         elif value.startswith("chart"):
#             print(f"starts with chart: {value}")
#             charts[key] = value.replace("chart-", "")
#         else:
#             print(f"starts with image: {value}")
#             images[key] = value
#     return icons, charts, images


import os
import xml.etree.ElementTree as ET
from zipfile import ZipFile
import re
from enum import Enum

def parse_images_with_links_and_captions(docx_path):
    """
    Extracts each image instance from a .docx file with:
    - image file path
    - hyperlink (if any)
    - image metadata (descr / name)
    - associated figure caption and number
    - inferred numbering scheme
    """
    image_info = []
    
    with ZipFile(docx_path) as docx_zip:
        document_xml = ET.fromstring(docx_zip.read("word/document.xml"))
        rels_xml = ET.fromstring(docx_zip.read("word/_rels/document.xml.rels"))

    NS = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }
    RELS_NS = '{http://schemas.openxmlformats.org/package/2006/relationships}'

    # Map rId → Target (image or link)
    rels_lookup = {
        rel.attrib['Id']: rel.attrib['Target']
        for rel in rels_xml.findall(f"{RELS_NS}Relationship")
    }

    # Get all paragraphs (for matching captions)
    paragraphs = document_xml.findall('.//w:p', NS)

    for i, paragraph in enumerate(paragraphs):
        drawing = paragraph.find('.//w:drawing', NS)
        if drawing is not None:
            # Metadata
            docpr = drawing.find('.//wp:docPr', NS)
            descr = docpr.attrib.get("descr") if docpr is not None else ""
            name = docpr.attrib.get("name") if docpr is not None else ""

            # Image
            blip = drawing.find('.//a:blip', NS)
            img_rid = blip.attrib.get(f"{{{NS['r']}}}embed") if blip is not None else None
            img_target = rels_lookup.get(img_rid)

            # Hyperlink
            hlink = drawing.find('.//a:hlinkClick', NS)
            link_rid = hlink.attrib.get(f"{{{NS['r']}}}id") if hlink is not None else None
            link_target = rels_lookup.get(link_rid)

            # Try to get caption from the NEXT paragraph
            caption_para = paragraphs[i + 1] if i + 1 < len(paragraphs) else None
            fig_label, fig_number, caption_text, numbering_type, fig_caption = None, None, "", "", ""

            if caption_para is not None:
                # Look for caption text and figure numbering
                texts = caption_para.findall('.//w:t', NS)
                fld = caption_para.find('.//w:fldSimple', NS)
                instr = fld.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr') if fld is not None else ""

                if 'SEQ Figure' in instr:
                    match = re.search(r'SEQ Figure(\\\* [A-Z]+)?', instr)
                    if match:
                        if '\\* ARABIC' in instr:
                            numbering_type = "numeric"
                        elif '\\* ALPHABETIC' in instr:
                            numbering_type = "alphabetic"
                        elif '\\* ROMAN' in instr:
                            numbering_type = "roman"
                        else:
                            numbering_type = "unknown"

                    # Get actual number from <w:t> inside fldSimple
                    fig_number_elem = fld.find('.//w:t', NS)
                    fig_number = fig_number_elem.text if fig_number_elem is not None else None

                    # Get caption text from rest of paragraph
                    caption_text = ''.join(t.text for t in texts).strip()

                    # Try to isolate "Figure X" as label
                    fig_label = re.match(r'Figure\s+\S+', caption_text)
                    fig_label = fig_label.group(0) if fig_label else f"Figure {fig_number}"

                    # Extract prefix before "Figure"
                    # Extract everything before "Figure X" (with flexible matching)
                    if fig_number:
                        # Updated pattern to match "Figure" followed by an optional letter and then capture the caption
                        pattern = r'^(Figure(?: [A-Z])?)(\s*)(.*)'
                        match = re.match(pattern, caption_text)
                        if match:
                            fig_prefix, space, fig_suffix = match.groups()
                            fig_label = fig_prefix + space
                            fig_caption = fig_suffix[len(fig_number):] ## remove the number from caption text

            image_info.append({
                "image_id": img_rid,
                "image_file": f"./{img_target}",
                "link_id": link_rid,
                "link_url": link_target,
                "alt_text": descr or name or "",
                "figure_number": fig_number,
                "figure_label": fig_label,
                "figure_caption": fig_caption,
                "caption_text": caption_text,
                "numbering_type": numbering_type
            })

    return image_info


# Create new figure captions
def update_figure_numbers(keep_image_map):
    """
    Update figure numbers in the image map based on new postion.
    For example if Figures 1-3 are removed, then Figure 4 becomes Figure 1.
    """
    figure_counter = {}
    for image in keep_image_map:
        if image["figure_number"] is not None:
            # Extract the figure number from the caption text
            figure_prefix = image["figure_label"]
            if figure_prefix not in figure_counter:
                figure_counter[figure_prefix] = 1
            else:
                figure_counter[figure_prefix] += 1
            # Update the figure number and caption text
            image["figure_number_new"] = str(figure_counter[figure_prefix])
            image["caption_text_new"] = ''.join([image["figure_label"], image["figure_number_new"], image["figure_caption"]])
        else:
            image["figure_number_new"] = None
            
    return keep_image_map


# remove keep from alt text
def identify_image_type(keep_image_map):
    for keep_image in keep_image_map:
        _, image_type, *image_alt_text = keep_image["alt_text"].split("-")
        image_alt_text = "-".join(image_alt_text)
        print(image_type)
        if image_type == "icon":
            keep_image["image_type"] = image_type
            ## Add additional tags
            icon_class, icon_name = image_alt_text.split(":")
            keep_image["class"] = icon_class
            keep_image["icon"] = icon_name
            ##TODO come up with another way of storing icon caption
            ## check if there is a caption - Nowhere else to store caption text because the risk icons include the caption in the image
            # if len(alt_text_split) > 2:
            #     keep_image["caption"] = alt_text_keep_removed.split("-")[2]
        elif image_type == "chart":
            keep_image["image_type"] = image_type
            ## Add additional tags
            keep_image["chart"] = image_alt_text
        elif image_type == "image":
            keep_image["image_type"] = image_type
            ## Add additional tags
            keep_image["alt_text_new"] = image_alt_text

    return keep_image_map

def replace_images_with_divs(html_content, image_mapping):

    class IconType(Enum):
        DROUGHT = "{{droughtRisk}}"
        FLOOD = "{{floodRisk}}"
        BUSHFIRE = "{{bushfireRisk}}"


    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Process all <img> tags with the source matches the image dictionary key in image_mapping
    for i, img in enumerate(soup.find_all('img')):
        image_meta = image_mapping[i]
        src = img.get('src')
        alt_text = img.get('alt')

        if src == image_meta["image_file"] and alt_text == image_meta["alt_text"]:
            print(image_meta)
            if image_meta["image_type"] == "icon": ## Caption for the risk icon can be pulled in dynamically in content-snippets.json
                print("icon")
                # Replace <img> tag with <div> tag with data-icon attribute
                div_tag = soup.new_tag("div")
                div_tag['data-icon-name'] = image_meta["icon"]
                print(image_meta["icon"])
                div_tag['data-icon-type'] = image_meta["class"].split("-")[0]
                if image_meta["class"] != "custom":
                    div_tag['data-style'] = image_meta["class"].split("-")[2]
                if image_meta["icon"] == "drought":
                    div_tag['data-caption'] = IconType.DROUGHT.value
                elif image_meta["icon"] == "flood":
                    div_tag['data-caption'] = IconType.FLOOD.value
                elif image_meta["icon"] == "bushfire":
                    div_tag['data-caption'] = IconType.BUSHFIRE.value
                div_tag.string = f""
                div_tag['class'] = "icon"
                if image_meta["link_url"] != "" and image_meta["link_url"] is not None:
                    div_tag['data-link'] = image_meta["link_url"]
                img.replace_with(div_tag)

            elif image_meta["image_type"] == "chart":
                print("chart")
                if src == image_meta["image_file"] and alt_text == image_meta["alt_text"]:
                    # Replace <img> tag with <div> tag with data-icon attribute
                    div_tag = soup.new_tag("div")
                    div_tag['id'] = image_meta["chart"] ## chart is a unique identifier for the chart hence used as id
                    div_tag['data-caption'] = image_meta["caption_text_new"]
                    div_tag.string = f""
                    div_tag['class'] = "chart"
                    if image_meta["link_url"] != "" and image_meta["link_url"] is not None:
                        div_tag['data-link'] = image_meta["link_url"]
                    img.replace_with(div_tag)

            elif image_meta["image_type"] == "image":
                print("image")
                # Replace <img> tag with <div> tag with data-icon attribute
                div_tag = soup.new_tag("div")
                div_tag['data-src'] = f"./assets/{image_meta["alt_text_new"]}.png" ## chart is a unique identifier for the chart hence used as id
                div_tag['data-caption'] = image_meta["caption_text_new"]
                div_tag.string = f""
                div_tag['class'] = "image"
                if image_meta["link_url"] != "" and image_meta["link_url"] is not None:
                    div_tag['data-link'] = image_meta["link_url"]
                img.replace_with(div_tag)

            print(f"Image {i} replaced with <div> tag with data-icon attribute: {div_tag}")
        else:
            print(f"Image {i} does not match the criteria for replacement.")
    # Get updated HTML as string
    return str(soup)

def replace_icons_with_placholders(html_content, icons_src: dict):
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Process all <img> tags with the source matches the image dictionary key in image_mapping
    for img in soup.find_all('img'):
        src = img.get('src')
        if src in icons_src:
            src_name = icons_src[src]
            print(src_name)
            if src_name:
                icon_div = soup.new_tag("div")
                icon_div['data-icon'] = src_name
                icon_div.string = f""
                icon_div['class'] = "icon"
                img.replace_with(icon_div)

    # Get updated HTML as string
    html_content = str(soup)
    return html_content

def replace_charts_with_placeholders(html_content, charts_src: dict, figure_captions: dict):
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Normalize chart src keys by basename
    normalized_charts_src = {os.path.basename(k): v for k, v in charts_src.items()}
    print(f"Normalized chart src: {normalized_charts_src}")

    # Process all <figure> tags containing <img> tags with the source matching the charts_src keys
    for figure in soup.find_all('figure'):
        img = figure.find('img')
        if img:
            src = img.get('src')
            src_basename = os.path.basename(src)

            if src_basename in normalized_charts_src:
                src_name = normalized_charts_src[src_basename]
                if src_name:
                    chart_div = soup.new_tag("div")

                    # Lookup caption by image basename (e.g. 'image11')
                    key = os.path.splitext(src_basename)[0]
                    print("looking up caption for key: ", key)
                    if key in figure_captions:
                        chart_div['data-caption'] = f"Figure {figure_captions[key].figure_number}{figure_captions[key].figure_caption}"
                    # else:
                    #     chart_div['caption'] = "No caption available"

                    chart_div['id'] = f"{src_name}"
                    
                    chart_div.string = f"{src_name.capitalize()}"
                    chart_div['class'] = "chart"

                    figure.replace_with(chart_div)

    # Return updated HTML
    return str(soup)


def replace_images_with_placeholders(html_content, images_src: dict, figure_captions: dict):
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Normalize images src keys by basename
    normalized_images_src = {os.path.basename(k): v for k, v in images_src.items()}
    print(f"Normalized images src: {normalized_images_src}")

    # Process all <figure> tags containing <img> tags with the source matching the images_src keys
    for figure in soup.find_all('figure'):
        img = figure.find('img')
        if img:
            src = img.get('src')
            src_basename = os.path.basename(src)

            if src_basename in normalized_images_src:
                src_name = normalized_images_src[src_basename]
                if src_name:
                    image_div = soup.new_tag("div")

                    # Lookup caption by image basename (e.g. 'image11')
                    key = os.path.splitext(src_basename)[0]
                    print("looking up caption for key: ", key)
                    if key in figure_captions:
                        image_div['data-caption'] = f"Figure {figure_captions[key].figure_number}{figure_captions[key].figure_caption}"
                    # else:
                    #     image_div['caption'] = "No caption available"

                    image_div['class'] = "image"
                    image_div['data-src'] = f"assets/{src_basename}"
                    image_div.string = f"{src_name.capitalize()}"

                    figure.replace_with(image_div)

    # Return updated HTML
    return str(soup)

# def replace_images_with_placeholders(html_content, alt_text_map):
#     """
#     Replace images with placeholders using regex and store alt text & figcaptions in alt_text_map.

#     Args:
#         html_content (str): The HTML content containing images.
#         alt_text_map (dict): A dictionary mapping image filenames to placeholder values.

#     Returns:
#         tuple: Modified HTML content (str) and updated alt_text_map (dict).
#     """

#     # Regex to match <figure> elements containing <img> and <figcaption>
#     figure_pattern = re.compile(
#         r'<figure[^>]*>\s*(<img[^>]*src=["\']([^"\']+)["\'][^>]*alt=["\']([^"\']*)["\'][^>]*>)\s*(<figcaption[^>]*>(.*?)</figcaption>)?\s*</figure>',
#         re.IGNORECASE | re.DOTALL
#     )

#     def replacer(match, figure_counter = [0]):      # mutable figure_counter so it can be incremented inside replacer()
#         img_src = match.group(2)
#         alt_text = match.group(3) if match.group(3) else ""
#         figcaption_text = match.group(5).strip() if match.group(5) else ""
#         img_id = os.path.splitext(os.path.basename(img_src))[0]

#         if figcaption_text:
#             print(f"Processing caption for image {img_id}")
#             figure_counter[0] += 1
#             figure_number = figure_counter[0]
#             # Only modify if image is in alt_text_map
#             if img_id not in alt_text_map:
#                 print(f"Image {img_id} not found in alt_text_map. Skipping replacement.")
#                 return match.group(0)  # return original <figure> block unchanged

#             if not figcaption_text.startswith("<p>Figure "):
#                 if figcaption_text.startswith("<p>"):
#                     figcaption_text = figcaption_text[3:]
#                 print(f"'Figure ' not found in caption text for figure {figure_number}. Adding manually as: <p>Figure {figure_number}{figcaption_text}")
#                 figcaption_text = f"<p>Figure {figure_number}{figcaption_text}"
#         else:
#             figure_number = None  # No caption, so no number needed

#         # Store alt text and caption
#         alt_text_map[img_id]['alt_text'] = alt_text.removeprefix(ALT_TEXT_KEEP_PREFIX) if alt_text else ""
#         alt_text_map[img_id]['figcaption'] = figcaption_text

#         # Escape caption
#         figcaption_escaped = html.escape(figcaption_text)

#         # Insert placeholder
#         placeholder_html = f'<div data-image="{alt_text_map[img_id]["path"]}" caption="{figcaption_escaped}">{alt_text_map[img_id]["alt_text"]}</div><br>'
#         return placeholder_html

#     # Apply regex replacement
#     modified_html = re.sub(figure_pattern, replacer, html_content)

#     return modified_html, alt_text_map


# def replace_icons_with_placeholders(html_content, alt_text_map):
#     """
#     Replace images with placeholders using regex and store alt text & figcaptions in alt_text_map.

#     Args:
#         html_content (str): The HTML content containing images.
#         alt_text_map (dict): A dictionary mapping image filenames to placeholder values.

#     Returns:
#         tuple: Modified HTML content (str) and updated alt_text_map (dict).
#     """

#     # Regex to match <figure> elements containing <img> and <figcaption>
#     pattern = re.compile(
#         r'<img[^>]*src=["\']([^"\']+)["\'][^>]*alt=["\']([^"\']*)["\'][^>]*\/?>',
#         re.IGNORECASE | re.DOTALL
#     )


#     def replacer(match, icon_counter=[0]):
#         print(f"Processing icon {icon_counter[0]}")
#         print(f"Matching alt text: {match.group(2)}")
#         img_src = match.group(1)
#         alt_text = match.group(2) if match.group(2) else ""
#         alt_text = alt_text.replace("keep-", "")

#         # Get ID or filename from img_src (e.g., "image2.png")
#         img_id = img_src.split("/")[-1]
#         print(f"Image ID: {img_id}")
#         # remove the file extension
#         img_id = os.path.splitext(img_id)[0]
#         print(f"Image ID split: {img_id}")

#         # Escape caption (in case it's needed)
#         figcaption_escaped = alt_text_map[img_id]["alt_text"].replace('"', '&quot;')

#         # Construct replacement div
#         placeholder_html = (
#             f'<div data-icon="{alt_text}">'
#             f'{alt_text_map[img_id]["alt_text"]}</div>'
#         )
#         return placeholder_html


#     # Apply regex replacement
#     modified_html = re.sub(pattern, replacer, html_content)

#     return modified_html, alt_text_map


def convert_docx_to_html(doc_path: str, lua_script: str,  keep_images: list):
    """
    Converts a DOCX file to HTML, removes image tags, and embeds custom CSS for Poppins font.

    :param doc_path: Path to the DOCX file.
    :param output_path: Path to save the output HTML file.
    """
    # Media to keep formatted for the lua script used by pypandoc
    metadata_json = generate_lua_lookup_table(keep_images)

    # Convert DOCX to HTML
    html = pypandoc.convert_file(
        doc_path, 
        "html", 
        extra_args=[
            "--quiet",
            f"--lua-filter={lua_script}",  # Replace with your actual Lua filter file
            "--extract-media=.",  # Extract media to the current directory
            # "--metadata", f"keep_images={metadata_json}"  # Pass as JSON
            "--metadata", f"keep_images={json.dumps(keep_images)}"  # Pass as JSON
        ]
    )

    return html

