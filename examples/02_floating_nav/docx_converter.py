from docx import Document
from lxml import etree
import json
from docx.shared import Pt, RGBColor
from collections import Counter
from enum import Enum
import pypandoc
import re
from bs4 import BeautifulSoup
import os
import zipfile
import shutil
import xml.etree.ElementTree as ET

# Define the prefix for alt text that should be kept in the output
ALT_TEXT_KEEP_PREFIX = "keep-"

## Define the folder structure for the output
FOLDERS = {
    "content": "content",
    "media": "assets",
    "data": "data",
    "js": "js",
    "css": "css",
}

# Define system-wide defaults as a simple dictionary
GLOBAL_DEFAULT_STYLES = {
    "textAlign": "left",
    "fontSize": "1rem",  # 16px converted to rem
    "fontWeight": "normal",
    "color": "#000000",
    "backgroundColor": "#FFFFFF",
    "padding": "8px",  # 5px converted to rem
    "verticalAlign": "middle"
}

# Function to map Word vertical alignment values to CSS
def map_vertical_align(word_val: str) -> str:
    """Convert Word's vertical alignment values to CSS equivalents."""
    mapping = {
        "top": "top",
        "center": "middle",  # Word uses "center", CSS uses "middle"
        "bottom": "bottom",
        "both": "middle"  # No direct equivalent, "middle" is a safe choice
    }
    return mapping.get(word_val.lower(), "middle")  # Default to "middle"


def convert_pt_to_rem(pt_size, base_font_size=16):
    """
    Converts font sizes from a DOCX file (in pt) to rem for HTML rendering.
    :param pt_size: Font size in points.
    :param base_font_size: Base font size in pixels (default: 16px).
    :return: Dictionary mapping text to its font size in rem.
    """
    rem_size = (pt_size * 1.3333) / base_font_size

    return f"{rem_size:.2f}rem"

def get_font_info(style):
    """Extract font information from a paragraph or run style."""
    font_info = {}

    if style and style.font:
        if style.font.name:
            font_info["fontFamily"] = style.font.name
        if style.font.size:
            font_info["fontSize"] = convert_pt_to_rem(style.font.size.pt)
        if style.font.bold:
            font_info["fontWeight"] = "bold"
        if style.font.italic:
            font_info["fontStyle"] = "italic"
        if style.font.color and style.font.color.rgb:
            font_info["color"] = f"#{style.font.color.rgb}"  # Convert RGBColor to hex

    return font_info

## TABLES
# Default styles for tables
DEFAULT_STYLES = {
        "table": {

        },
        "th": {

        },
        "td1": {

        },
        "td": {

        }
    }


def get_table_alignment(table):
    tbl_pr = table._element.xpath(".//w:jc")
    if tbl_pr:
        return tbl_pr[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
    return "default"

def get_cell_vertical_alignment(cell):
    v_align = cell._element.xpath(".//w:vAlign")
    if v_align:
        return v_align[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
    return "center"

def get_paragraph_alignment(paragraph):
    p_pr = paragraph._element.xpath(".//w:jc")
    if p_pr:
        return p_pr[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
    return "left"

def clean_dict(data):
    """
    Recursively removes keys with values of '#auto' or None from a dictionary.
    """
    if isinstance(data, dict):
        return {k: clean_dict(v) for k, v in data.items() if v not in ('#auto', None)}
    elif isinstance(data, list):
        return [clean_dict(v) for v in data]
    else:
        return data


# Function to compute the modal values while ignoring missing properties
def compute_modal_style(styles_list):
    # Add the default styles to ensure they are considered in the modal calculation
    styles_list.append(GLOBAL_DEFAULT_STYLES.copy())

    modal_styles = {}
    all_keys = {key for styles in styles_list for key in styles}
    
    for key in all_keys:
        values = [styles[key] for styles in styles_list if key in styles]
        if values and Counter(values).most_common(1)[0][0] is not None:
            modal_styles[key] = Counter(values).most_common(1)[0][0]

    return modal_styles

def extract_table_default_styles(doc) -> dict:
    table_styles = {
        "table": {
            "border": "1px solid black",
            "borderCollapse": "collapse",
            "marginBottom": "12px"
        },
        "th": {},
        "td1": {},
        "td": {}
    }

    if not doc.tables:
        return table_styles

    table = doc.tables[1]  # Assume second table defines the styles
    default_font_size_rem = convert_pt_to_rem(get_doc_default_font_size(doc)) ## Default font size extraction

    th_styles = []
    td1_styles = []
    td_styles = []


    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_style = {}

            # Extract text alignment
            text_align = None
            for para in cell.paragraphs:
                if para.text.strip():
                    text_align = get_paragraph_alignment(para)
                    break
            if text_align:
                cell_style["textAlign"] = text_align

            # Extract font styles
            font_size = None
            font_weight = None
            font_color = None

            for para in cell.paragraphs:
                if para.text.strip():
                    for run in para.runs:
                        if run.font.size:
                            font_size = convert_pt_to_rem(run.font.size.pt)
                        elif para.style and para.style.font.size:
                            font_size = convert_pt_to_rem(para.style.font.size.pt)

                        if run.bold:
                            font_weight = "bold"

                        if run.font.color and run.font.color.rgb:
                            font_color = f"#{run.font.color.rgb}"
                        elif para.style and para.style.font.color and para.style.font.color.rgb:
                            font_color = f"#{para.style.font.color.rgb}"
                    break  # Only need the first valid paragraph with text

            # if font_size:
            #     cell_style["fontSize"] = font_size
            # if font_weight:
            #     cell_style["fontWeight"] = font_weight
            # if font_color:
            #     cell_style["color"] = font_color
            cell_style["fontSize"] = font_size if font_size else default_font_size_rem
            cell_style["fontWeight"] = font_weight if font_weight else None
            cell_style["color"] = font_color if font_color else None

            # Extract background color
            shading = cell._element.xpath('.//w:shd/@w:fill')
            cell_style["backgroundColor"] = f"#{shading[0]}" if shading else None  # Leave None to filter later
            
            # Padding and vertical alignment
            cell_style["padding"] = "5px"
            cell_style["verticalAlign"] = map_vertical_align(get_cell_vertical_alignment(cell))

            # Store styles in respective lists
            if row_idx == 0:
                th_styles.append(cell_style)
            elif col_idx == 0:
                td1_styles.append(cell_style)
            else:
                td_styles.append(cell_style)

    # Compute the modal styles for each category
    table_styles["th"] = compute_modal_style(th_styles)
    table_styles["td1"] = compute_modal_style(td1_styles)
    table_styles["td"] = compute_modal_style(td_styles)

    return clean_dict(table_styles)


## END TABLES

## Default font size extraction
def get_style_font_size(doc, style_name):
    """Fetch font size from a named style if available."""
    try:
        style = doc.styles[style_name]
        if style and style.font and style.font.size:
            return style.font.size.pt
    except KeyError:
        pass
    return None  # Style not found or no explicit size set



def get_doc_default_font_size(doc):
    """Extract the default font size from the document."""
    # Get the XML element as a string
    xml_data = doc.styles.element.xml

    # Define the XML namespace
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Parse the XML
    root = ET.fromstring(xml_data)

    # Find the <w:docDefaults> section and locate <w:sz>
    sz_element = root.find('.//w:docDefaults//w:sz', ns)

    # Extract the font size
    if sz_element is not None:
        font_size = sz_element.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        return int(font_size) / 2 # Convert half-points to points 
    return None


def extract_styles(doc_path):
    """Extract styles from the Word document dynamically."""
    doc = Document(doc_path)

    styles_data = {
        "headings": {},
        "body": {},
        "lists": {},
        "captions": {},
    }

    for para in doc.paragraphs:
        para_style = para.style

        if para_style and para_style.name.startswith("Heading"):
            heading_level = para_style.name.replace("Heading ", "h")
            styles_data["headings"][heading_level] = get_font_info(para_style)

        elif para_style and "Caption" in para_style.name:
            styles_data["captions"]["caption"] = get_font_info(para_style)

        elif para_style and "List" in para_style.name:
            list_type = "ul" if "Bullet" in para_style.name else "ol"
            styles_data["lists"][list_type] = get_font_info(para_style)

        elif para_style and para_style.name == "Normal":
            styles_data["body"]["p"] = get_font_info(para_style)

    # Extract default font size
    default_font_size = get_style_font_size(doc, "Normal") or get_doc_default_font_size(doc) or 12
    styles_data["body"]["p"]["fontSize"] = convert_pt_to_rem(default_font_size)

    ## TABLES
    if doc.tables:
        table = extract_table_default_styles(doc)
        styles_data.update(table)
    return styles_data

def is_cell_merged(cell, row_idx, col_idx, merge_tracker):
    """Check if a Word table cell is merged (horizontally or vertically) and determine row span."""
    cell_xml = cell._tc  # Get the XML element of the table cell

    # Check for horizontal merge (gridSpan)
    grid_span = cell_xml.xpath('.//w:gridSpan')
    col_span = int(grid_span[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", 1)) if grid_span else 1

    # Check for vertical merge (vMerge)
    v_merge = cell_xml.xpath('.//w:vMerge')
    row_span = None  # Only set if it's actually merged

    if v_merge:
        v_merge_val = v_merge[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
        if v_merge_val == "restart":  # Start of a vertically merged section
            merge_tracker[(row_idx, col_idx)] = 1  # Initialize tracking
        elif v_merge_val is None:  # Continuation of merge
            for r in range(row_idx - 1, -1, -1):  # Find the start of this vertical merge
                if (r, col_idx) in merge_tracker:
                    merge_tracker[(r, col_idx)] += 1
                    row_span = merge_tracker[(r, col_idx)]
                    return {"hidden": True}  # Mark this cell as part of the merged block, but not stored separately

    return {
        "colSpan": col_span if col_span > 1 else None,  # Store only if >1
        "rowSpan": row_span if row_span and row_span > 1 else None  # Store only if >1
    }




from docx import Document

def extract_table_format(doc_path, default_styles: dict = DEFAULT_STYLES):
    doc = Document(doc_path)
    tables_info = {}

    for table_idx, table in enumerate(doc.tables):
        table_id = f"table_{table_idx}"  # Generate table ID
        table_info = {
            "headers": [],
            "rows": []
        }

        merge_tracker = {}  # Track merged cells {(row_idx, col_idx): remaining_span}

        for row_idx, row in enumerate(table.rows):
            row_data = []
            col_idx = 0  # Track actual column position for skipping merged cells

            while col_idx < len(row.cells):
                if (row_idx, col_idx) in merge_tracker:
                    merge_tracker[(row_idx, col_idx)] -= 1
                    if merge_tracker[(row_idx, col_idx)] == 0:
                        del merge_tracker[(row_idx, col_idx)]  # Clear when done
                    col_idx += 1
                    continue  # Skip merged cells

                cell = row.cells[col_idx]

                # Determine default style based on position
                if row_idx == 0:
                    default_style = DEFAULT_STYLES["th"]
                elif col_idx == 0:
                    default_style = DEFAULT_STYLES["td1"]
                else:
                    default_style = DEFAULT_STYLES["td"]

                # Extract formatting and text as separate parts
                text_parts = []
                for para in cell.paragraphs:
                    if para.text.strip():
                        for run in para.runs:
                            part = {"text": run.text}
                            if run.bold:
                                part["bold"] = True
                            if run.italic:
                                part["italic"] = True
                            if run.font.superscript:
                                part["superscript"] = True
                            if run.font.subscript:
                                part["subscript"] = True
                            if run.font.color and run.font.color.rgb:
                                part["color"] = f"#{run.font.color.rgb}"
                            if run.font.size:
                                part["fontSize"] = convert_pt_to_rem(run.font.size.pt)
                            text_parts.append(part)

                # Default to single text if no formatting is found
                if len(text_parts) == 1:
                    actual_style = {
                        "text": text_parts[0]["text"].replace("\n", "<br>"),
                    }
                else:
                    actual_style = {
                        "textParts": text_parts
                    }

                # Handle merged cells
                merge_info = is_cell_merged(cell, row_idx, col_idx, merge_tracker)
                if "hidden" in merge_info:
                    col_idx += 1
                    continue  # Skip storing this cell (it's a continuation of a merged cell)

                if merge_info["colSpan"]:
                    actual_style["colSpan"] = merge_info["colSpan"]
                    for i in range(1, merge_info["colSpan"]):  # Skip following columns
                        merge_tracker[(row_idx, col_idx + i)] = 1

                if merge_info["rowSpan"]:
                    actual_style["rowSpan"] = merge_info["rowSpan"]
                    for i in range(1, merge_info["rowSpan"]):  # Track vertically merged cells
                        merge_tracker[(row_idx + i, col_idx)] = merge_info["rowSpan"] - i

                # Extract other styles
                actual_style["textAlign"] = get_paragraph_alignment(cell.paragraphs[0]) if cell.paragraphs else None
                actual_style["verticalAlign"] = map_vertical_align(get_cell_vertical_alignment(cell))

                # Extract background color
                shading = cell._element.xpath('.//w:shd/@w:fill')
                actual_style["backgroundColor"] = f"#{shading[0]}" if shading else None

                # Remove default styles
                actual_style = clean_dict(actual_style)

                # Determine if this is a header row or data row
                if row_idx == 0:
                    table_info["headers"].append(actual_style)
                else:
                    row_data.append(actual_style)

                col_idx += 1  # Move to the next column

            if row_idx > 0:  # Store only data rows (headers handled separately)
                table_info["rows"].append(row_data)

        tables_info[table_id] = table_info

    return tables_info




def extract_docx(doc_path, output_path):
    """Extracts the DOCX contents into a specified folder."""
    extracted_folder = os.path.join(output_path, "docx_extracted")
    os.makedirs(extracted_folder, exist_ok=True)
    
    with zipfile.ZipFile(doc_path, "r") as docx_zip:
        docx_zip.extractall(extracted_folder)
    
    return extracted_folder

def parse_relationships(rels_path):
    """Parses the relationships file to map image IDs to filenames."""
    image_map = {}
    
    if os.path.exists(rels_path):
        tree = ET.parse(rels_path)
        root = tree.getroot()
        
        for rel in root.findall(".//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"):
            rid = rel.attrib.get("Id", "")
            target = rel.attrib.get("Target", "")
            
            if "media/" in target:
                image_map[rid] = target.split("/")[-1]
    
    return image_map

def extract_alt_texts(doc_xml_path, image_map, allowed_alt_texts, extracted_folder, output_path, image_folder):
    """Extracts images based on allowed alt texts and renames them."""
    alt_text_map = {}

    # Create the output folder for images
    media_folder = os.path.join(output_path, image_folder)
    os.makedirs(image_folder, exist_ok=True)
    
    if os.path.exists(doc_xml_path):
        tree = ET.parse(doc_xml_path)
        root = tree.getroot()

        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main", 
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        }

        for drawing in root.findall(".//w:drawing", ns):
            doc_pr = drawing.find(".//a:blip", ns)
            descr_tag = drawing.find(".//wp:docPr", ns)
            
            if doc_pr is not None and descr_tag is not None:
                alt_text = descr_tag.attrib.get("descr", "").strip()
                rid = doc_pr.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed", "")
                
                if rid in image_map and (alt_text in allowed_alt_texts or alt_text.startswith(ALT_TEXT_KEEP_PREFIX)):
                    old_name = image_map[rid]
                    alt_text_map[os.path.splitext(old_name)[0]] = f"{image_folder}/{old_name}"

                    old_path = os.path.join(extracted_folder, "word/media", old_name)
                    new_path = os.path.join(media_folder, old_name)

                    if os.path.exists(old_path):
                        print(f"✅ Moving image: {old_name} ➝ {new_path}")
                        shutil.move(old_path, new_path)
                    else:
                        print(f"❌ ERROR: Image file not found: {old_path}")
    
    return alt_text_map


def extract_docx_media(doc_path, output_path, media_folder, allowed_alt_texts):
    """Extracts DOCX contents and images based on allowed alt texts."""
    # Extract DOCX contents
    extracted_folder = extract_docx(doc_path, output_path)

    # Paths to XML files inside the extracted DOCX
    rels_path = os.path.join(extracted_folder, "word/_rels/document.xml.rels")
    doc_xml_path = os.path.join(extracted_folder, "word/document.xml")

    # Parse relationships to map image IDs to filenames
    image_map = parse_relationships(rels_path)

    alt_text_map = extract_alt_texts(doc_xml_path, image_map, allowed_alt_texts, extracted_folder, output_path, media_folder)

    return alt_text_map

def generate_lua_lookup_table(image_numbers):
    """
    Generates a Pandoc-compatible metadata JSON for Lua.
    Example: [3, 5, 6] -> { "keep_images": [3, 5, 6] }
    """
    metadata = {
        "keep_images": image_numbers  # Ensure it's a list, not a string
    }

    # Convert metadata to JSON
    metadata_json = json.dumps(metadata)

    return metadata_json

# Replace tables sequentially with placeholders
def table_replacer(match, counter=[0]):
    replacement = f'<div data-table="table_{counter[0]}"></div>'
    counter[0] += 1
    return replacement


def generate_navigation_data(html_content):
    """
    Parses HTML content to generate structured navigation data.

    :param html_content: HTML content as a string.
    :return: Structured navigation data as a list of dictionaries.
    """
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Organize headings into a structured hierarchy
    nav_data = []
    current_h1 = None
    current_h2 = None

    for tag in soup.find_all(['h1', 'h2', 'h3']):
        heading_id = tag.get('id')
        if not heading_id:
            heading_id = re.sub(r'\s+', '-', tag.text.strip().lower())
            tag['id'] = heading_id

        if tag.name == 'h1':
            current_h1 = {
                "id": heading_id,
                "text": tag.text,
                "h2": []
            }
            nav_data.append(current_h1)

        elif tag.name == 'h2' and current_h1:
            current_h2 = {
                "id": heading_id,
                "text": tag.text,
                "h3": []
            }
            current_h1["h2"].append(current_h2)

        elif tag.name == 'h3' and current_h1 and current_h2:
            current_h2["h3"].append({
                "id": heading_id,
                "text": tag.text
            })

    return nav_data

def insert_sub_navigation(html_content, nav_json):
    """
    Inserts <div data-sub-navigation data-parent="section_id"></div> before the first <h2> of each section.
    """
    for section in nav_json:
        section_id = section["id"]
        h2_items = section.get("h2", [])
        
        # Only proceed if there are H2 elements in the section
        if not h2_items:
            continue
        
        # Define the sub-navigation div to be inserted
        sub_nav_div = f'\n<div data-sub-navigation data-parent="{section_id}"></div>'
        
        # Find the first H2 in this section and insert the sub-navigation div before it
        pattern = rf'(\n<h2 id="{h2_items[0]["id"]}")'
        html_content = re.sub(pattern, sub_nav_div + r'\1', html_content, count=1)
    
    return html_content

## Image html replacement
def replace_images_with_placeholders(html_content, alt_text_map):
    """
    Replace images with placeholders using regex and store alt text in alt_text_map.

    Args:
        html_content (str): The HTML content containing images.
        alt_text_map (dict): A dictionary mapping image filenames to placeholder values.

    Returns:
        tuple: Modified HTML content (str) and updated alt_text_map (dict).
    """

    # Updated regex to match <img> tags regardless of attribute order
    img_pattern = re.compile(
        r'<img[^>]*src=["\']([^"\']+)["\'][^>]*alt=["\']([^"\']*)["\'][^>]*>',
        re.IGNORECASE
    )

    def replacer(match):
        img_src = match.group(1)  # Extract src
        #print(img_src)
        alt_text = match.group(2) if match.group(2) else ""  # Extract alt (if exists)
        #print(alt_text)
        img_id = os.path.splitext(os.path.basename(img_src))[0]  # Extract image ID

        # Store alt text if not already in map
        alt_text_map[img_id]['alt_text'] = alt_text.removeprefix(ALT_TEXT_KEEP_PREFIX) if alt_text else ""

        # Replace <img> with a <div> preserving the alt text or placeholder
        return f'<div data-image="{alt_text_map[img_id]['path']}">{alt_text_map[img_id]['alt_text']}</div><br>'

    # Apply regex replacement
    modified_html = re.sub(img_pattern, replacer, html_content)

    return modified_html, alt_text_map





def convert_docx_to_html(doc_path: str, lua_script: str,  keep_images: list):
    """
    Converts a DOCX file to HTML, removes image tags, and embeds custom CSS for Poppins font.

    :param doc_path: Path to the DOCX file.
    :param output_path: Path to save the output HTML file.
    """
    # Media to keep formatted for the lua script used by pypandoc
    metadata_json = generate_lua_lookup_table(keep_images)

    # Convert DOCX to HTML
    html = pypandoc.convert_file(
        doc_path, 
        "html", 
        extra_args=[
            "--quiet",
            f"--lua-filter={lua_script}",  # Replace with your actual Lua filter file
            "--metadata", f"keep_images={metadata_json}"  # Pass as JSON
        ]
    )

    return html

